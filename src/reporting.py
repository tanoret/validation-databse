from __future__ import annotations

import io
import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from reportlab.lib import pagesizes
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    PageBreak,
)
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn


DEFAULT_CRITICAL_CHARACTERISTICS = [
    {
        "id": "CC-01",
        "name": "Functional correctness (verification)",
        "description": "Demonstrates the code solves the intended equations and numerical methods correctly for the targeted features.",
        "evidence_types": ["analytic benchmarks", "manufactured solutions", "regression tests", "code-to-code"],
    },
    {
        "id": "CC-02",
        "name": "Accuracy vs data (validation)",
        "description": "Demonstrates agreement with experimental/benchmark data within defined uncertainties for the intended domain.",
        "evidence_types": ["SET/IET validation", "uncertainty-aware comparisons"],
    },
    {
        "id": "CC-03",
        "name": "Applicability / intended-use coverage",
        "description": "Connects evidence to the user’s intended use and defines applicability boundaries (envelope of validity).",
        "evidence_types": ["use-case demonstrations", "range-of-applicability mapping"],
    },
    {
        "id": "CC-04",
        "name": "Reproducibility & configuration control",
        "description": "Shows controlled, repeatable execution with version/commit, run recipe, and traceability to inputs/outputs.",
        "evidence_types": ["build/installation checks", "run recipes", "hashes/artifact integrity"],
    },
    {
        "id": "CC-05",
        "name": "Problem reporting & corrective action",
        "description": "Known issues are documented; impacts assessed; fixes traceable to evidence re-runs (as applicable).",
        "evidence_types": ["issue log", "impact assessment", "re-run after fix"],
    },
]


def _now_iso() -> str:
    return datetime.utcnow().replace(microsecond=0).isoformat() + "Z"


def _md_escape(s: str) -> str:
    return (s or "").replace("\n", " ").strip()


def _mk_case_row(case: Dict[str, Any], reports: Dict[str, Any]) -> Dict[str, Any]:
    # Build a compact row for tables
    sr = case.get("source_reports", []) or []
    report_refs: List[str] = []
    for item in sr:
        rid = item.get("report_id")
        rep = reports.get(rid, {}) if rid else {}
        num = rep.get("report_number") or rid or "Unknown report"
        note = item.get("note") or ""
        report_refs.append(f"{num} — {note}".strip(" —"))
    return {
        "id": case.get("id"),
        "title": case.get("title"),
        "vv_type": case.get("vv_type"),
        "scope": case.get("scope"),
        "tools": ", ".join(case.get("tools", []) or []),
        "system": case.get("system"),
        "reports": "; ".join(report_refs),
    }



def _p(styles, text: str, style_name: str = "Normal") -> Paragraph:
    """Safe paragraph for table cells (wraps long text)."""
    s = (text or "").replace("\n", " ").strip()
    return Paragraph(s, styles[style_name])


def _make_case_catalog_table(selected_cases: List[Dict[str, Any]], reports: Dict[str, Any], styles, doc_width: float) -> Table:
    """Create a compact, wrapped table that fits the page."""
    header = ["Case ID", "Title", "V&V type", "Scope", "Tools", "Source reports"]
    data = [[_p(styles, h, "BodyText") for h in header]]

    for c in selected_cases:
        row = _mk_case_row(c, reports)
        data.append([
            _p(styles, str(row["id"])),
            _p(styles, str(row["title"])),
            _p(styles, str(row["vv_type"])),
            _p(styles, str(row["scope"])),
            _p(styles, str(row["tools"])),
            _p(styles, str(row["reports"])),
        ])

    # Column widths as fractions of available doc width
    fracs = [0.14, 0.28, 0.10, 0.10, 0.14, 0.24]  # sum = 1.0
    col_widths = [doc_width * f for f in fracs]

    tbl = Table(data, colWidths=col_widths, repeatRows=1, splitByRow=1)
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#EDEDED")),
        ("GRID", (0,0), (-1,-1), 0.25, colors.grey),
        ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
        ("FONTSIZE", (0,0), (-1,-1), 8),
        ("LEADING", (0,0), (-1,-1), 9),
        ("VALIGN", (0,0), (-1,-1), "TOP"),
        ("LEFTPADDING", (0,0), (-1,-1), 4),
        ("RIGHTPADDING", (0,0), (-1,-1), 4),
        ("TOPPADDING", (0,0), (-1,-1), 3),
        ("BOTTOMPADDING", (0,0), (-1,-1), 3),
    ]))
    return tbl


def _set_docx_table_borders(table):
    """Add simple borders to a python-docx table."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '4')
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), 'A0A0A0')
        tblBorders.append(element)
    tblPr.append(tblBorders)


def generate_cgd_markdown(
    db: Dict[str, Any],
    selected_case_ids: List[str],
    user_inputs: Dict[str, Any],
    critical_characteristics: Optional[List[Dict[str, Any]]] = None,
) -> str:
    cases = db.get("cases", {})
    reports = db.get("reports", {})
    critical_characteristics = critical_characteristics or DEFAULT_CRITICAL_CHARACTERISTICS

    selected_cases: List[Dict[str, Any]] = []
    for cid in selected_case_ids:
        c = cases.get(cid)
        if c:
            cc = dict(c)
            cc.setdefault("id", cid)
            selected_cases.append(cc)

    title = user_inputs.get("package_title") or "Commercial Grade Dedication Evidence Package (Draft)"
    tool_name = user_inputs.get("tool_name") or "NEAMS Tool(s)"
    tool_version = user_inputs.get("tool_version") or "TBD"
    intended_use = user_inputs.get("intended_use") or "TBD"
    environment = user_inputs.get("execution_environment") or "TBD"
    requester = user_inputs.get("requester") or "TBD"
    reviewer = user_inputs.get("reviewer") or "TBD"

    md: List[str] = []
    md.append(f"# { _md_escape(title) }\n")
    md.append(f"**Generated:** {_now_iso()}  ")
    md.append(f"**Prepared for:** {_md_escape(requester)}  ")
    md.append(f"**Reviewer/Approver:** {_md_escape(reviewer)}\n")

    md.append("## 1. Purpose and scope\n")
    md.append(f"This document is a **draft CGD-ready evidence package** for **{_md_escape(tool_name)}** (version **{_md_escape(tool_version)}**).  ")
    md.append("It consolidates verification/validation/benchmark/demonstration evidence from the NEAMS-MSRs Validation Database and provides traceability to the originating milestone reports.\n")
    md.append(f"**Intended use:** {_md_escape(intended_use)}\n")

    md.append("## 2. Software identification and configuration\n")
    md.append(f"- **Tool(s):** {_md_escape(tool_name)}\n- **Version/Commit:** {_md_escape(tool_version)}\n- **Execution environment:** {_md_escape(environment)}\n")
    md.append("(Add compiler, libraries, platform, container hash, and run recipes as available.)\n")

    md.append("## 3. Critical characteristics and evidence strategy\n")
    for cc in critical_characteristics:
        md.append(f"- **{cc['id']} – {cc['name']}**: {cc['description']}")
    md.append("\n")

    md.append("## 4. Evidence summary (selected cases)\n")
    md.append(f"Total selected cases: **{len(selected_cases)}**\n")

    md.append("### 4.1 Case catalog\n")
    md.append("| Case ID | Title | V&V type | Scope | Tools | Source reports |\n|---|---|---:|---:|---|---|")

    for c in selected_cases:
        row = _mk_case_row(c, reports)
        md.append(
            f"| `{_md_escape(row['id'])}` | {_md_escape(row['title'])} | {_md_escape(row['vv_type'])} | {_md_escape(row['scope'])} | {_md_escape(row['tools'])} | {_md_escape(row['reports'])} |"
        )
    md.append("\n")

    md.append("### 4.2 Case details (appendix-style)\n")
    for c in selected_cases:
        md.append(f"#### {_md_escape(c.get('id'))}: {_md_escape(c.get('title'))}\n")
        md.append(f"- **V&V type:** {_md_escape(c.get('vv_type'))}")
        md.append(f"- **Scope:** {_md_escape(c.get('scope'))}")
        md.append(f"- **System:** {_md_escape(c.get('system'))}")
        md.append(f"- **Tools:** {', '.join(c.get('tools', []) or [])}")
        md.append(f"- **Phenomena:** {', '.join(c.get('phenomena', []) or [])}")
        md.append(f"- **Summary:** {_md_escape(c.get('summary',''))}\n")
        # Metrics
        metrics = c.get("metrics", []) or []
        if metrics:
            md.append("**Metrics:**")
            for m in metrics:
                name = m.get("name")
                val = m.get("value")
                basis = m.get("basis")
                md.append(f"- {name}: {val} ({basis})")
        # References to reports
        md.append("**Source report traceability:**")
        for sr in c.get("source_reports", []) or []:
            rid = sr.get("report_id")
            rep = reports.get(rid, {}) if rid else {}
            num = rep.get("report_number") or rid or "Unknown report"
            rep_title = rep.get("title") or ""
            rep_link = rep.get("file_link") or sr.get("report_link") or ""
            note = sr.get("note") or ""
            sec = sr.get("section") or ""
            md.append(f"- {num} — {rep_title}  ")
            if rep_link:
                md.append(f"  - link: {rep_link}")
            if sec:
                md.append(f"  - section: {sec}")
            if note:
                md.append(f"  - note: {note}")
        md.append("\n")

    md.append("## 5. Limitations and required reviewer actions\n")
    md.append("- This package is automatically generated and requires independent technical review.\n")
    md.append("- Add explicit acceptance criteria and uncertainty treatment for each validation comparison as required by the dedication plan.\n")
    md.append("- Ensure the final package includes controlled input decks, run scripts, and output artifacts (checksummed).\n")

    return "\n".join(md).strip() + "\n"


def render_cgd_pdf(
    db: Dict[str, Any],
    selected_case_ids: List[str],
    user_inputs: Dict[str, Any],
    out_path: Path,
) -> Path:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    cases = db.get("cases", {})
    reports = db.get("reports", {})

    selected_cases: List[Dict[str, Any]] = []
    for cid in selected_case_ids:
        c = cases.get(cid)
        if c:
            cc = dict(c)
            cc.setdefault("id", cid)
            selected_cases.append(cc)

    title = user_inputs.get("package_title") or "Commercial Grade Dedication Evidence Package (Draft)"
    tool_name = user_inputs.get("tool_name") or "NEAMS Tool(s)"
    tool_version = user_inputs.get("tool_version") or "TBD"
    intended_use = user_inputs.get("intended_use") or "TBD"
    environment = user_inputs.get("execution_environment") or "TBD"

    styles = getSampleStyleSheet()
    # Tighten body text for denser tables
    styles['BodyText'].fontSize = 8
    styles['BodyText'].leading = 9
    styles['Normal'].fontSize = 10
    styles['Normal'].leading = 12
    story = []

    story.append(Paragraph(title, styles["Title"]))
    story.append(Spacer(1, 0.2 * inch))

    story.append(Paragraph(f"Generated: {_now_iso()}", styles["Normal"]))
    story.append(Paragraph(f"Tool(s): {tool_name}", styles["Normal"]))
    story.append(Paragraph(f"Version/Commit: {tool_version}", styles["Normal"]))
    story.append(Paragraph(f"Execution environment: {environment}", styles["Normal"]))
    story.append(Spacer(1, 0.2 * inch))

    story.append(Paragraph("Purpose and scope", styles["Heading2"]))
    story.append(Paragraph(f"Intended use: {intended_use}", styles["Normal"]))
    story.append(Spacer(1, 0.2 * inch))


    story.append(Paragraph("Selected case catalog", styles["Heading2"]))

    # Use wrapped table sized to doc width to avoid overflow
    # Note: doc.width is available after creating doc; compute with LETTER minus margins used below.
    page_w, page_h = pagesizes.LETTER
    left_margin = 0.7 * inch
    right_margin = 0.7 * inch
    doc_width = page_w - left_margin - right_margin

    tbl = _make_case_catalog_table(selected_cases, reports, styles, doc_width)
    story.append(tbl)
    story.append(PageBreak())

    story.append(Paragraph("Case details", styles["Heading2"]))
    for c in selected_cases:
        story.append(Paragraph(f"{c.get('id')}: {c.get('title')}", styles["Heading3"]))
        story.append(Paragraph(f"V&V type: {c.get('vv_type')} — Scope: {c.get('scope')} — System: {c.get('system')}", styles["Normal"]))
        story.append(Paragraph(f"Tools: {', '.join(c.get('tools', []) or [])}", styles["Normal"]))
        story.append(Paragraph(f"Phenomena: {', '.join(c.get('phenomena', []) or [])}", styles["Normal"]))
        if c.get("summary"):
            story.append(Paragraph(f"Summary: {c.get('summary')}", styles["Normal"]))
        # Metrics
        metrics = c.get("metrics", []) or []
        if metrics:
            story.append(Paragraph("Metrics:", styles["Normal"]))
            for m in metrics:
                story.append(Paragraph(f"- {m.get('name')}: {m.get('value')} ({m.get('basis','')})", styles["Normal"]))
        # Traceability
        story.append(Paragraph("Source report traceability:", styles["Normal"]))
        for sr in c.get("source_reports", []) or []:
            rid = sr.get("report_id")
            rep = reports.get(rid, {}) if rid else {}
            num = rep.get("report_number") or rid or "Unknown report"
            rep_title = rep.get("title") or ""
            note = sr.get("note") or ""
            story.append(Paragraph(f"- {num}: {rep_title} — {note}", styles["Normal"]))
        story.append(Spacer(1, 0.15 * inch))

    doc = SimpleDocTemplate(str(out_path), pagesize=pagesizes.LETTER, leftMargin=0.7*inch, rightMargin=0.7*inch, topMargin=0.75*inch, bottomMargin=0.75*inch)
    doc.build(story)
    return out_path


def render_cgd_docx(
    db: Dict[str, Any],
    selected_case_ids: List[str],
    user_inputs: Dict[str, Any],
    out_path: Path,
) -> Path:
    """Create a CGD package as a DOCX (Word) document."""
    out_path.parent.mkdir(parents=True, exist_ok=True)
    cases = db.get("cases", {})
    reports = db.get("reports", {})

    selected_cases: List[Dict[str, Any]] = []
    for cid in selected_case_ids:
        c = cases.get(cid)
        if c:
            cc = dict(c)
            cc.setdefault("id", cid)
            selected_cases.append(cc)

    title = user_inputs.get("package_title") or "Commercial Grade Dedication Evidence Package (Draft)"
    tool_name = user_inputs.get("tool_name") or "NEAMS Tool(s)"
    tool_version = user_inputs.get("tool_version") or "TBD"
    intended_use = user_inputs.get("intended_use") or "TBD"
    environment = user_inputs.get("execution_environment") or "TBD"
    requester = user_inputs.get("requester") or "TBD"
    reviewer = user_inputs.get("reviewer") or "TBD"

    doc = Document()

    # Title
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta_run = meta.add_run(f"Generated: {_now_iso()}\nPrepared for: {requester}\nReviewer/Approver: {reviewer}")
    meta_run.font.size = Pt(10)

    doc.add_paragraph()

    doc.add_heading("1. Purpose and scope", level=2)
    doc.add_paragraph(
        f"This document is a draft CGD-ready evidence package for {tool_name} (version/commit: {tool_version}). "
        "It consolidates V&V/benchmark/demonstration evidence and preserves traceability to source reports."
    )
    doc.add_paragraph(f"Intended use: {intended_use}")

    doc.add_heading("2. Software identification and configuration", level=2)
    doc.add_paragraph(f"Tool(s): {tool_name}")
    doc.add_paragraph(f"Version/Commit: {tool_version}")
    doc.add_paragraph(f"Execution environment: {environment}")

    doc.add_heading("3. Selected case catalog", level=2)
    table = doc.add_table(rows=1, cols=6)
    hdr = table.rows[0].cells
    headers = ["Case ID", "Title", "V&V type", "Scope", "Tools", "Source reports"]
    for i, h in enumerate(headers):
        hdr[i].text = h

    for c in selected_cases:
        row = _mk_case_row(c, reports)
        cells = table.add_row().cells
        cells[0].text = str(row["id"] or "")
        cells[1].text = str(row["title"] or "")
        cells[2].text = str(row["vv_type"] or "")
        cells[3].text = str(row["scope"] or "")
        cells[4].text = str(row["tools"] or "")
        cells[5].text = str(row["reports"] or "")

    _set_docx_table_borders(table)

    doc.add_page_break()
    doc.add_heading("4. Case details", level=2)

    for c in selected_cases:
        doc.add_heading(f"{c.get('id')}: {c.get('title')}", level=3)
        doc.add_paragraph(f"V&V type: {c.get('vv_type')} | Scope: {c.get('scope')} | System: {c.get('system')}")
        doc.add_paragraph(f"Tools: {', '.join(c.get('tools', []) or [])}")
        doc.add_paragraph(f"Phenomena: {', '.join(c.get('phenomena', []) or [])}")
        if c.get("summary"):
            doc.add_paragraph(f"Summary: {c.get('summary')}")

        metrics = c.get("metrics", []) or []
        if metrics:
            doc.add_paragraph("Metrics:")
            for m in metrics:
                doc.add_paragraph(f"- {m.get('name')}: {m.get('value')} ({m.get('basis','')})", style=None)

        doc.add_paragraph("Source report traceability:")
        for sr in c.get("source_reports", []) or []:
            rid = sr.get("report_id")
            rep = reports.get(rid, {}) if rid else {}
            num = rep.get("report_number") or rid or "Unknown report"
            rep_title = rep.get("title") or ""
            note = sr.get("note") or ""
            doc.add_paragraph(f"- {num}: {rep_title} — {note}")

    doc.save(str(out_path))
    return out_path
